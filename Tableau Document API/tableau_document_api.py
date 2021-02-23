#install tableaudocumentapoi for python

#pip install tableaudocumentapi

import tableaudocumentapi as td

#load the twbx file in to metadata
wb = td.Workbook(r'C:\Users\sivkumar\Documents\Project\Deloitte\Consulting_ExCo.twbx')

#pandas for dataframe calculations
import pandas as pd

#for creating word docx
import docx
#initiate a document
doc = docx.Document()
#add document Title
doc.add_heading('Tableau Documentation',level=1)

####################################################################
#Collecting data for data sources
ds = []
for i in range(len(wb.datasources)):
    ds.append(wb.datasources[i].caption)
#add Sub Title
doc.add_heading('Data source list', level=3)
#data in to dataframe
df_ds = pd.DataFrame(ds, columns=['Data Source Name'])
#make a table based on shape of data fame
t = doc.add_table(df_ds.shape[0]+1, df_ds.shape[1])
t.style = 'Table Grid'

for j in range(df_ds.shape[-1]):
    t.cell(0,j).text = df_ds.columns[j]
for i in range(df_ds.shape[0]):
    for j in range(df_ds.shape[-1]):
        t.cell(i+1,j).text = str(df_ds.values[i,j])
        
####################################################################
#add Sub Title
doc.add_heading('Worksheet list', level=3)
#data in to dataframe
df_ws = pd.DataFrame(wb.worksheets, columns=['worksheet_name'])
#make a table based on shape of data fame
t = doc.add_table(df_ws.shape[0]+1, df_ws.shape[1])
t.style = 'Table Grid'

for j in range(df_ws.shape[-1]):
    t.cell(0,j).text = df_ws.columns[j]
for i in range(df_ws.shape[0]):
    for j in range(df_ws.shape[-1]):
        t.cell(i+1,j).text = str(df_ws.values[i,j])

####################################################################
#collect calculated feild data
field_id = []
field_name = []
fomula = []
for i in wb.datasources[1].fields.keys():
   field_id.append(wb.datasources[1].fields[i].id)
   field_name.append(wb.datasources[1].fields[i].caption)
   fomula.append(wb.datasources[1].fields[i].calculation)

#make it into data frame
df = pd.DataFrame.from_dict({'field_id' : field_id,
                            'field_name' : field_name,
                            'formula' : fomula})

#add Sub Title
doc.add_heading('Feild Data for Data source {0}'.format(wb.datasources[1].caption), level=3)
#df.shape
t = doc.add_table(df.shape[0]+1, df.shape[1])
t.style = 'Table Grid'

for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i+1,j).text = str(df.values[i,j])

####################################################################

doc.save(r'C:\Users\sivkumar\Documents\Github Repositories\Sample Data\tableau_documentation.docx')
